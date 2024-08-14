# data_parser.py
import re
from abc import ABC, abstractmethod
from docx import Document
import pymupdf  # PyMuPDF


class Parser(ABC):
    @abstractmethod
    def parse(self, file_path: str) -> list:
        pass


class ProcedureParser(Parser):
    def parse(self, file_path: str) -> list:
        if file_path.endswith(".docx"):
            document = Document(file_path)
            return self.chunk_procedures(document, file_path)
        elif file_path.endswith(".pdf"):
            return self.parse_pdf(file_path)
        else:
            raise ValueError(
                "Unsupported file format. Please upload a .docx or .pdf file."
            )

    def parse_pdf(self, file_path: str, chunk_size: int = 100) -> list:
        chunks = []
        current_chunk = []
        current_plain_chunk = []
        current_length = 0
        section_heading = None
        section_heading_plain = None

        doc = pymupdf.open(file_path)
        for page_num in range(doc.page_count):
            page = doc.load_page(page_num)
            text = page.get_text("text").strip()
            if text:
                paragraphs = text.split("\n")
                for para in paragraphs:
                    para = para.strip()
                    if para:
                        words = para.split()
                        if current_length + len(words) > chunk_size:
                            self.add_chunk(
                                chunks,
                                current_chunk,
                                current_plain_chunk,
                                section_heading,
                                section_heading_plain,
                                file_path,
                            )
                            current_chunk = [para]
                            current_plain_chunk = [para]
                            current_length = len(words)
                        else:
                            current_chunk.append(para)
                            current_plain_chunk.append(para)
                            current_length += len(words)

        if current_chunk:
            self.add_chunk(
                chunks,
                current_chunk,
                current_plain_chunk,
                section_heading,
                section_heading_plain,
                file_path,
            )

        return chunks

    def chunk_procedures(
        self, document: Document, file_path: str, chunk_size: int = 100
    ) -> list:
        chunks = []
        current_chunk = []
        current_plain_chunk = []
        current_length = 0
        section_heading = None
        section_heading_plain = None

        for para in document.paragraphs:
            text = para.text.strip()
            if text:
                plain_text = text
                formatted_text = self.format_paragraph(para)
                if para.style.name.startswith("Heading"):
                    if current_chunk:
                        self.add_chunk(
                            chunks,
                            current_chunk,
                            current_plain_chunk,
                            section_heading,
                            section_heading_plain,
                            file_path,
                        )
                        current_chunk = []
                        current_plain_chunk = []
                        current_length = 0
                    section_heading = formatted_text
                    section_heading_plain = plain_text
                else:
                    words = text.split()
                    if current_length + len(words) > chunk_size:
                        self.add_chunk(
                            chunks,
                            current_chunk,
                            current_plain_chunk,
                            section_heading,
                            section_heading_plain,
                            file_path,
                        )
                        current_chunk = [formatted_text]
                        current_plain_chunk = [plain_text]
                        current_length = len(words)
                    else:
                        current_chunk.append(formatted_text)
                        current_plain_chunk.append(plain_text)
                        current_length += len(words)

        if current_chunk:
            self.add_chunk(
                chunks,
                current_chunk,
                current_plain_chunk,
                section_heading,
                section_heading_plain,
                file_path,
            )

        # Now handle tables
        for table in document.tables:
            table_text = self.format_table(table)
            chunks.append(
                {
                    "filename": file_path,
                    "heading": None,
                    "plain_text": table_text,
                    "formatted_text": table_text,
                }
            )

        return chunks

    def add_chunk(
        self,
        chunks,
        current_chunk,
        current_plain_chunk,
        section_heading,
        section_heading_plain,
        file_path,
    ):
        chunk_text = " ".join(current_chunk)
        chunk_plain_text = " ".join(current_plain_chunk)
        if section_heading:
            chunk_text = f"{section_heading}\n{chunk_text}"
            chunk_plain_text = f"{section_heading_plain}\n{chunk_plain_text}"
        chunks.append(
            {
                "filename": file_path,
                "heading": section_heading,
                "plain_text": chunk_plain_text,
                "formatted_text": chunk_text,
            }
        )

    def format_paragraph(self, para):
        runs = []
        for run in para.runs:
            text = run.text
            if run.bold:
                text = f"<strong>{text}</strong>"
            if run.italic:
                text = f"<em>{text}</em>"
            if run.underline:
                text = f"<u>{text}</u>"
            if run.font.color.rgb:
                color = run.font.color.rgb
                text = f'<span style="color: rgb({color[0]}, {color[1]}, {color[2]});">{text}</span>'
            runs.append(text)

        formatted_text = "".join(runs)
        if para.style.name.startswith("Heading"):
            level = int(para.style.name[-1])
            formatted_text = f"<h{level}>{formatted_text}</h{level}>"
        elif para.style.name == "List Paragraph":
            formatted_text = f"<li>{formatted_text}</li>"

        return formatted_text

    def format_table(self, table):
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(cells)
        table_html = "<table>"
        for row in rows:
            table_html += "<tr>" + "".join(f"<td>{cell}</td>" for cell in row) + "</tr>"
        table_html += "</table>"
        return table_html
